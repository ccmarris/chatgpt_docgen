import logging
import docx

_logger = logging.getLogger(__name__)

def save_responses(prompt_response_pairs, filename, output_format):
    '''
    Save the generated responses to a file in the specified format.
    Parameters:
        prompt_response_pairs (list): A list of tuples containing prompts and their corresponding responses.
        filename (str): The name of the output file.
        output_format (str): The format to save the responses in ("docx", "txt", or "md").
    Returns:
        None
    '''
    if output_format == "docx":
        _logger.info(f"Saving responses to {filename} in DOCX format")
        # Use the docgen module to save responses to a Word document
        save_responses_to_docx(prompt_response_pairs, filename=filename)
    elif output_format == "txt":
        _logger.info(f"Saving responses to {filename} in TXT format")
        # Save responses to a text file
        with open(filename, "w", encoding="utf-8") as f:
            for prompt, response in prompt_response_pairs:
                f.write(f"Prompt: {prompt}\nResponse: {response}\n\n")
    elif output_format == "md":
        _logger.info(f"Saving responses to {filename} in Markdown format")
        # Save responses to a Markdown file
        with open(filename, "w", encoding="utf-8") as f:
            for prompt, response in prompt_response_pairs:
                f.write(f"## Prompt\n{prompt}\n\n### Response\n{response}\n\n")
    elif output_format == "stdout":
        _logger.info("Printing responses to stdout")
        # Print responses to standard output
        for prompt, response in prompt_response_pairs:
            print(f"## Prompt: {prompt}\n## Response: {response}\n")
    else:
        _logger.error(f"Unknown output format: {output_format}")
    
    return


def save_responses_to_docx(prompt_response_pairs, filename="output.docx"):
    doc = docx.Document()
    for prompt, response in prompt_response_pairs:
        doc.add_heading(prompt, level=2)
        doc.add_paragraph(response)
    doc.save(filename)